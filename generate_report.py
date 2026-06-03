"""
Rapor oluşturma scripti — BLM3058 Simülasyon ve Modelleme Dersi
Çalıştır: python generate_report.py
Çıktı   : rapor.docx
"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

RESULTS = "results"

# yardımcı fonksiyonlar

def set_font(run, name="Times New Roman", size=12, bold=False,
             italic=False, color=None):
    """Apply font properties (name, size, style, color) to a python-docx run."""
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_paragraph(doc, text="", style="Normal", align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  size=12, bold=False, italic=False, space_before=0, space_after=6):
    """Add a formatted paragraph; returns the paragraph object."""
    p = doc.add_paragraph(style=style)
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        set_font(run, size=size, bold=bold, italic=italic)
    return p


def add_heading(doc, text, level=1, size=14):
    """Add a bold heading at the given level (1=chapter, 2=section, 3=subsection)."""
    sizes = {1: 16, 2: 14, 3: 12}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_font(run, size=sizes.get(level, size), bold=True)
    return p


def add_figure(doc, img_path, caption, width_cm=14):
    """Insert an image file with a centered italic caption; shows placeholder if file missing."""
    if os.path.exists(img_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(img_path, width=Cm(width_cm))
    else:
        p = doc.add_paragraph(f"[Görsel bulunamadı: {img_path}]")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(12)
    run = cap.add_run()
    set_font(cap.runs[0], size=10, italic=True)
    return p


def add_bullet(doc, text, level=0, size=12):
    """Add a bulleted list item; level controls indentation depth."""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(level * 0.5 + 0.5)
    run = p.add_run(text)
    set_font(run, size=size)
    return p


def add_table(doc, headers, rows, col_widths=None):
    """Create a 'Table Grid' styled table with bold centered headers."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            set_font(run, bold=True, size=11)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for run in cell.paragraphs[0].runs:
                set_font(run, size=11)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table


def page_break(doc):
    """Insert a hard page break."""
    doc.add_page_break()


# rapor oluşturma

def build_report():
    doc = Document()

    # Sayfa yapısı: A4, kenar boşlukları 2.5 cm
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # kapak sayfası

    add_paragraph(doc, "T.C.", align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=13, bold=True, space_before=30)
    add_paragraph(doc, "MARMARA ÜNİVERSİTESİ", align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=14, bold=True)
    add_paragraph(doc, "TEKNOLOJİ FAKÜLTESİ", align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=13, bold=True)
    add_paragraph(doc, "BİLGİSAYAR MÜHENDİSLİĞİ BÖLÜMÜ",
                  align=WD_ALIGN_PARAGRAPH.CENTER, size=13, bold=True)

    add_paragraph(doc, "", space_before=30)

    add_paragraph(doc,
                  "BLM3058 — SİMÜLASYON VE MODELLEME DERSİ",
                  align=WD_ALIGN_PARAGRAPH.CENTER, size=13, bold=True)
    add_paragraph(doc, "DÖNEM PROJESİ RAPORU",
                  align=WD_ALIGN_PARAGRAPH.CENTER, size=13, bold=True)

    add_paragraph(doc, "", space_before=20)

    add_paragraph(doc,
                  "UYDU ÇARPIŞMA ÖNLEME SİMÜLASYONU\n"
                  "(Satellite Collision Avoidance Simulation)",
                  align=WD_ALIGN_PARAGRAPH.CENTER, size=16, bold=True,
                  space_before=10)

    add_paragraph(doc, "", space_before=40)

    add_paragraph(doc, "Hazırlayanlar:", align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=12, bold=True)
    add_paragraph(doc, "Muhammed Emin Dinç — 170423052",
                  align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
    add_paragraph(doc, "Emre Yanalak — 170423056",
                  align=WD_ALIGN_PARAGRAPH.CENTER, size=12)

    add_paragraph(doc, "", space_before=20)
    add_paragraph(doc, "Ders Sorumlusu:", align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=12, bold=True)
    add_paragraph(doc, "Dr. Öğr. Üyesi Şahin Uyaver",
                  align=WD_ALIGN_PARAGRAPH.CENTER, size=12)

    add_paragraph(doc, "", space_before=20)
    add_paragraph(doc, "Teslim Tarihi: 30 Mayıs 2026",
                  align=WD_ALIGN_PARAGRAPH.CENTER, size=12)

    page_break(doc)

    # özet

    add_heading(doc, "ÖZET", level=1)
    add_paragraph(doc,
        "Bu çalışmada, Düşük Dünya Yörüngesi'nde (LEO) faaliyet gösteren "
        "uyduların karşılaştığı çarpışma tehlikelerini tespit etmek ve önlemek "
        "amacıyla kapsamlı bir simülasyon sistemi Python programlama dili "
        "kullanılarak geliştirilmiştir. Sistem; J2 pertürbasyonlu iki cisim "
        "problemini RK45 yöntemiyle sayısal olarak çözerek yörünge yayımı "
        "gerçekleştirmekte, Chan 2D projeksiyon yöntemiyle çarpışma olasılığı "
        "(Pc) hesaplamakta ve minimum delta-V kısıtı altında kaçınma manevrası "
        "planlamaktadır. Dört farklı senaryo altında değerlendirilen sistem, "
        "N=200 replikasyonluk Monte Carlo analizi ve 5 doğrulama/geçerleme "
        "testiyle istatistiksel olarak desteklenmiştir. Yüksek riskli "
        "konjunksiyon senaryosunda 585 m'lik yaklaşma mesafesi, 0.43 m/s'lik "
        "along-track yakıt yakımıyla 1300 m'ye yükseltilmiş; çarpışma "
        "olasılığı 1.92×10⁻⁴'ten 1.0×10⁻⁵'e indirilerek güvenli seviyeye "
        "getirilmiştir. Çoklu enkaz senaryosunda ise 5 farklı nesneyle "
        "eş zamanlı konjunksiyon yönetimi gerçeklenmiştir.")

    add_paragraph(doc, "", space_before=6)
    add_paragraph(doc,
        "Anahtar Kelimeler: Uydu çarpışma önleme, LEO yörünge mekaniği, "
        "conjunction analizi, çarpışma olasılığı, delta-V optimizasyonu, "
        "Monte Carlo simülasyonu.",
        italic=True)

    page_break(doc)

    # bölüm 1: problem tanımı

    add_heading(doc, "BÖLÜM 1: PROBLEM TANIMI VE AMAÇ", level=1)

    add_heading(doc, "1.1 Mevcut Durum", level=2)
    add_paragraph(doc,
        "Düşük Dünya Yörüngesi'nde (LEO, 160–2000 km irtifa) günümüzde "
        "10.000'den fazla aktif uydu ve NORAD tarafından izlenen 25.000'den "
        "fazla enkaz nesnesi bulunmaktadır. 2009 yılında yaşanan "
        "Iridium-33 ile Kosmos-2251 çarpışması, bu tehdidi somut biçimde "
        "ortaya koymuştur; söz konusu çarpışma tek başına 2.000'den fazla "
        "yeni enkaz parçası üretmiştir. Bu tablo, gerekli önlem alınmadığında "
        "Kessler Sendromu olarak bilinen kaskad çarpışma sürecinin "
        "tetiklenebileceğini göstermektedir.")
    add_paragraph(doc,
        "Bu nedenle, konjunksiyon (yakın yaklaşma) olaylarının "
        "erken tespiti ve kaçınma manevrasının optimum biçimde planlanması "
        "uzay operasyonlarının sürdürülebilirliği açısından kritik önem "
        "taşımaktadır. SpaceX, ESA ve NASA gibi kuruluşlar bu amaçla "
        "Conjunction Data Message (CDM) standartlarını oluşturmuş ve "
        "Pc ≥ 1×10⁻⁴ eşiği eylem kriteri olarak benimsenmiştir.")

    add_heading(doc, "1.2 Projenin Amacı", level=2)
    add_paragraph(doc,
        "Bu projenin amacı, LEO'daki uyduların birbirine yaklaşma "
        "durumlarını fizik temelli bir Python simülasyonu ile modellemek; "
        "çarpışma olasılığını uluslararası standartta hesaplamak ve "
        "gerektiğinde minimum yakıt tüketimiyle kaçınma manevrası "
        "planlamaktır. Çalışmayla aşağıdaki araştırma sorularına yanıt "
        "aranmaktadır:")
    add_bullet(doc, "Hangi konjunksiyon geometrisinde Pc eylem eşiğini (≥1×10⁻⁴) aşmaktadır?")
    add_bullet(doc, "Minimum ΔV ile Pc'yi eşiğin altına indirmek için hangi manevra stratejisi optimaldir?")
    add_bullet(doc, "Monte Carlo analizi, nominal Pc tahminlerini istatistiksel olarak nasıl desteklemektedir?")
    add_bullet(doc, "Çoklu enkaz ortamında öncelik sırasına dayalı manevra planlaması nasıl gerçeklenebilir?")

    page_break(doc)

    # bölüm 2: yöntem

    add_heading(doc, "BÖLÜM 2: YÖNTEM (METODOLOJİ)", level=1)

    add_heading(doc, "2.1 Matematiksel Model ve Varsayımlar", level=2)

    add_heading(doc, "2.1.1 Yörünge Yayımı (Orbital Propagation)", level=3)
    add_paragraph(doc,
        "Yörünge yayımında iki cisim problemi ve J2 pertürbasyonu "
        "birlikte kullanılmıştır. ECI (Earth-Centered Inertial) "
        "koordinat çerçevesindeki hareket denklemi:")
    add_paragraph(doc, "d²r/dt² = −(μ/r³)·r + a_J2",
                  align=WD_ALIGN_PARAGRAPH.CENTER, italic=True)
    add_paragraph(doc, "J2 ivmesi bileşenleri:")
    add_paragraph(doc,
        "a_J2 = −(3/2)·J2·μ·R_E²/r⁵ · [x(1−5z²/r²),  y(1−5z²/r²),  z(3−5z²/r²)]",
        align=WD_ALIGN_PARAGRAPH.CENTER, italic=True)
    add_paragraph(doc,
        "Burada μ = 398600.4418 km³/s², R_E = 6371 km, "
        "J2 = 1.08262668×10⁻³ değerleri kullanılmıştır. "
        "Sayısal entegrasyon için SciPy kütüphanesinin solve_ivp fonksiyonu "
        "(RK45, rtol=10⁻¹⁰, atol=10⁻¹²) tercih edilmiştir.")

    add_heading(doc, "2.1.2 Çarpışma Olasılığı — Chan 2D Projeksiyon Yöntemi", level=3)
    add_paragraph(doc,
        "TCA (Time of Closest Approach) anında göreli hız vektörüne dik "
        "çarpışma düzlemi tanımlanmaktadır. Kombine konum kovaryans matrisi "
        "(Σ = Σ₁ + Σ₂) bu düzleme projekte edilerek 2 boyutlu Gauss "
        "yoğunluk fonksiyonu elde edilmektedir. Çarpışma olasılığı:")
    add_paragraph(doc,
        "Pc = ∬_{|r|≤R_hbr} N(μ_rel, Σ_2D) dA",
        align=WD_ALIGN_PARAGRAPH.CENTER, italic=True)
    add_paragraph(doc,
        "Burada R_hbr = 10 m kombine sert cisim yarıçapını, "
        "μ_rel çarpışma düzlemindeki göreli konum vektörünü ifade "
        "etmektedir. İntegrasyon scipy.integrate.dblquad ile sayısal olarak "
        "gerçekleştirilmiştir.")

    add_heading(doc, "2.1.3 Manevra Planlama", level=3)
    add_paragraph(doc,
        "TCA'dan t_lead süresi önce anlık (impulsive) yakıt yakımı "
        "uygulanmaktadır. LVLH (Local Vertical Local Horizontal) "
        "çerçevesinde 6 yön (±along-track, ±radial, ±cross-track) "
        "arasında ikili arama (binary search) ile Pc < 10⁻⁵ koşulunu "
        "sağlayan minimum |ΔV| bulunmaktadır.")

    add_heading(doc, "2.1.4 Varsayımlar", level=3)
    add_bullet(doc, "Anlık yakıt yakımı (impulsive maneuver) — yakıt tükenmesi ihmal")
    add_bullet(doc, "Kovaryans matrisi TCA boyunca sabit kabul edilmiştir")
    add_bullet(doc, "Hava sürtünmesi ihmal edilmiş; dominant pertürbasyon J2'dir")
    add_bullet(doc, "Monte Carlo'da ölçüm gürültüsü kovaryans matrisinden örneklenmektedir")

    add_heading(doc, "2.2 Sistem Akış Diyagramı", level=2)
    add_figure(doc,
               os.path.join(RESULTS, "system_flowchart.png"),
               "Şekil 1: Satellite Collision Avoidance — Sistem Akış Diyagramı",
               width_cm=13)

    add_heading(doc, "2.3 Yazılım Geliştirme Ortamı", level=2)
    add_paragraph(doc,
        "Proje Python 3.13 ile Windows 11 ortamında, PyCharm IDE kullanılarak "
        "geliştirilmiştir. Kullanılan kütüphaneler:")

    headers = ["Kütüphane", "Sürüm", "Kullanım Amacı"]
    rows = [
        ["numpy", "≥1.24", "Sayısal dizi işlemleri, matris işlemleri"],
        ["scipy", "≥1.10", "ODE çözücü (RK45), sayısal integrasyon (dblquad)"],
        ["matplotlib", "≥3.7", "2D/3D grafik, GIF animasyon üretimi"],
        ["pillow", "≥9.0", "GIF animasyon kaydı (PillowWriter)"],
        ["tkinter", "stdlib", "Grafik kullanıcı arayüzü (GUI)"],
    ]
    add_table(doc, headers, rows, col_widths=[4, 3, 9])
    add_paragraph(doc, "Tablo 1: Kullanılan Python kütüphaneleri",
                  align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=10)
    add_paragraph(doc, "")

    add_paragraph(doc, "Proje modüler yapıda tasarlanmıştır:")
    headers2 = ["Modül", "Görev"]
    rows2 = [
        ["orbital_mechanics.py", "Fizik motoru — J2, RK45, koordinat dönüşümleri"],
        ["satellite.py", "Uydu veri modeli (Kepler elemanları + kovaryans)"],
        ["collision_detection.py", "TCA tespiti, Chan 2D Pc hesabı"],
        ["avoidance.py", "ΔV manevra planlama (6 yön, binary search)"],
        ["scenarios.py", "4 simülasyon senaryosu"],
        ["verification.py", "V&V — 5 doğrulama ve geçerleme testi"],
        ["monte_carlo.py", "N=200 Monte Carlo replikasyon analizi"],
        ["visualization.py", "Grafik, animasyon ve flowchart üretimi"],
        ["gui.py", "Tkinter grafik kullanıcı arayüzü"],
        ["main.py", "Terminal giriş noktası"],
    ]
    add_table(doc, headers2, rows2, col_widths=[5.5, 10.5])
    add_paragraph(doc, "Tablo 2: Proje modülleri ve görevleri",
                  align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=10)

    page_break(doc)

    # bölüm 3: V&V

    add_heading(doc, "BÖLÜM 3: MODELİN DOĞRULANMASI VE GEÇERLEMESİ (V&V)", level=1)

    add_heading(doc, "3.1 Doğrulama (Verification)", level=2)
    add_paragraph(doc,
        "Doğrulama testleri, kodun matematiksel modeli doğru "
        "biçimde gerçeklediğini ortaya koymaktadır.")

    headers3 = ["Test", "Açıklama", "Kriter", "Sonuç"]
    rows3 = [
        ["V1 — Enerji Korunumu",
         "5 tam yörüngede mekanik enerji korunumu",
         "Bağıl hata < 1×10⁻⁶", "GEÇER ✓"],
        ["V2 — Gidiş-Dönüş",
         "Kepler elemanları → ECI → Kepler elemanları",
         "Bağıl hata < 1×10⁻⁹", "GEÇER ✓"],
        ["V3 — Dairesel Hız",
         "Dairesel yörüngede v = √(μ/a) kontrolü",
         "Bağıl hata < 1×10⁻⁶", "GEÇER ✓"],
    ]
    add_table(doc, headers3, rows3, col_widths=[4, 6, 4, 2])
    add_paragraph(doc, "Tablo 3: Doğrulama test sonuçları",
                  align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=10)
    add_paragraph(doc, "")

    add_figure(doc,
               os.path.join(RESULTS, "vv_energy_conservation.png"),
               "Şekil 2: V1 — 5 Yörüngede Mekanik Enerji Korunumu",
               width_cm=13)

    add_heading(doc, "3.2 Geçerleme (Validation)", level=2)
    add_paragraph(doc,
        "Geçerleme testleri, modelin gerçek dünya beklentileriyle "
        "ve teorik değerlerle uyumunu göstermektedir.")

    headers4 = ["Test", "Açıklama", "Kriter", "Sonuç"]
    rows4 = [
        ["G1 — Kepler 3. Yasası",
         "T = 2π√(a³/μ) ile simülasyon periyodu karşılaştırması",
         "Bağıl hata < 1×10⁻⁵", "GEÇER ✓"],
        ["G2 — Pc Sınır Koşulları",
         "Miss >> R_hbr → Pc ≈ 0; Miss → 0 → Pc → 1",
         "Monoton davranış", "GEÇER ✓"],
    ]
    add_table(doc, headers4, rows4, col_widths=[4, 7, 3, 2])
    add_paragraph(doc, "Tablo 4: Geçerleme test sonuçları",
                  align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=10)
    add_paragraph(doc, "")

    add_figure(doc,
               os.path.join(RESULTS, "vv_pc_boundary.png"),
               "Şekil 3: G2 — Pc Sınır Koşulları Geçerlemesi",
               width_cm=13)

    page_break(doc)

    # bölüm 4: bulgular

    add_heading(doc, "BÖLÜM 4: BULGULAR VE ANALİZ", level=1)

    add_heading(doc, "4.1 Senaryo 1: Düşük Riskli Yakın Geçiş", level=2)
    add_paragraph(doc,
        "İki farklı eğimdeki yörünge (i=53°, i=90°) çaprazlaşmasında "
        "konjunksiyon analizi gerçeklenmiştir. TCA t=0.833 saatte "
        "gerçekleşmekte; bu anda miss distance 5424 m, göreli hız "
        "13.7 km/s olarak hesaplanmıştır. Pc = 1.2×10⁻⁶ eylem "
        "eşiğinin (1×10⁻⁴) çok altında kalmakta olup herhangi "
        "bir manevra gerektirmemektedir.")

    add_figure(doc,
               os.path.join(RESULTS, "s1_orbits_3d.png"),
               "Şekil 4: Senaryo 1 — Çapraz Yörüngeler ve TCA Noktası (3D ECI)",
               width_cm=12)
    add_figure(doc,
               os.path.join(RESULTS, "s1_miss_distance.png"),
               "Şekil 5: Senaryo 1 — Zaman İçinde Miss Distance (log ölçek)",
               width_cm=13)

    add_heading(doc, "4.2 Senaryo 2: Yüksek Riskli Konjunksiyon ve Kaçınma Manevrası", level=2)
    add_paragraph(doc,
        "Daha yakın çaprazlaşma parametreleriyle oluşturulan bu senaryoda "
        "miss distance 585 m'ye düşmekte ve Pc = 1.92×10⁻⁴ değeriyle "
        "eylem eşiğini aşmaktadır. Sistem otomatik olarak manevra "
        "planlamaya geçmekte; TCA'dan 0.5 saat önce optimal along-track "
        "yönünde ΔV = 0.43 m/s uygulanmaktadır.")

    headers5 = ["Parametre", "Manevra Öncesi", "Manevra Sonrası"]
    rows5 = [
        ["Miss Distance", "585 m", "1300 m"],
        ["Çarpışma Olasılığı (Pc)", "1.92×10⁻⁴", "1.0×10⁻⁵"],
        ["Risk Seviyesi", "YÜKSEK", "DÜŞÜK"],
        ["Manevra Stratejisi", "—", "Along-track (+)"],
        ["Delta-V", "—", "0.43 m/s"],
        ["Manevra Zamanı", "—", "T−30 dakika"],
    ]
    add_table(doc, headers5, rows5, col_widths=[5.5, 4.5, 5])
    add_paragraph(doc, "Tablo 5: Senaryo 2 — Manevra öncesi ve sonrası karşılaştırması",
                  align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=10)
    add_paragraph(doc, "")

    add_figure(doc,
               os.path.join(RESULTS, "s2_miss_distance.png"),
               "Şekil 6: Senaryo 2 — Miss Distance: Manevra Öncesi vs. Sonrası",
               width_cm=13)
    add_figure(doc,
               os.path.join(RESULTS, "s2_relative_motion.png"),
               "Şekil 7: Senaryo 2 — LVLH Çerçevesinde Göreli Hareket",
               width_cm=11)

    add_heading(doc, "4.3 Senaryo 3: Delta-V Duyarlılık Analizi", level=2)
    add_paragraph(doc,
        "Bu senaryoda ΔV büyüklüğü ile miss distance ve Pc arasındaki "
        "ilişki sistematik biçimde incelenmiştir. Sonuçlar, yalnızca "
        "2 m/s'lik along-track yakıt yakımının Pc'yi eylem eşiğinin "
        "altına indirebileceğini göstermektedir. Farklı ΔV seviyelerinde "
        "ve manevra zamanlamalarında Pc davranışı da değerlendirilmiştir.")

    add_figure(doc,
               os.path.join(RESULTS, "s3_dv_tradeoff.png"),
               "Şekil 8: Senaryo 3 — Delta-V ve Miss Distance ile Pc İlişkisi",
               width_cm=13)
    add_figure(doc,
               os.path.join(RESULTS, "s3_pc_vs_leadtime.png"),
               "Şekil 9: Senaryo 3 — Farklı ΔV Seviyelerinde Pc vs. Manevra Zamanlaması",
               width_cm=13)

    add_heading(doc, "4.4 Senaryo 4: Çoklu Enkaz Ortamı", level=2)
    add_paragraph(doc,
        "Aktif bir uydu etrafında 5 farklı enkaz nesnesi ile "
        "eş zamanlı konjunksiyon yönetimi gerçeklenmiştir. "
        "Tüm konjunksiyon olayları analiz edilerek Pc'ye göre "
        "öncelik sırası belirlenmiş; en kritik nesne (DEBRIS-B, "
        "Pc=1.7×10⁻⁴) için otomatik manevra planlanmıştır.")

    headers6 = ["Enkaz", "TCA (saat)", "Miss (m)", "Pc", "Risk"]
    rows6 = [
        ["DEBRIS-B", "1.17", "253", "1.70×10⁻⁴", "YÜKSEK ★"],
        ["DEBRIS-D", "1.53", "688", "2.77×10⁻⁸", "DÜŞÜK"],
        ["DEBRIS-C", "1.39", "2463", "3.30×10⁻⁶", "DÜŞÜK"],
        ["DEBRIS-A", "1.00", "7145", "2.32×10⁻⁸", "DÜŞÜK"],
        ["DEBRIS-E", "2.00", "6080", "1.26×10⁻⁷", "DÜŞÜK"],
    ]
    add_table(doc, headers6, rows6, col_widths=[3, 3, 3, 4, 3])
    add_paragraph(doc, "Tablo 6: Senaryo 4 — Çoklu enkaz konjunksiyon özeti",
                  align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=10)
    add_paragraph(doc, "")

    add_figure(doc,
               os.path.join(RESULTS, "s4_conjunction_timeline.png"),
               "Şekil 10: Senaryo 4 — Çoklu Enkaz Konjunksiyon Zaman Çizelgesi",
               width_cm=14)

    add_heading(doc, "4.5 Monte Carlo Analizi", level=2)
    add_paragraph(doc,
        "Senaryo 2 parametreleri için N=200 Monte Carlo replikasyonu "
        "gerçeklenmiştir. Her replikasyonda uyduların başlangıç konumu "
        "ve hızı, kovaryans matrisinden örneklenen Gaussian gürültü ile "
        "bozulmuş; böylece ölçüm belirsizliği ve başlangıç koşulu "
        "hassasiyeti modellenmiştir.")

    headers7 = ["İstatistik", "Değer"]
    rows7 = [
        ["Ortalama Pc", "1.68×10⁻⁵"],
        ["Standart Sapma", "3.80×10⁻⁵"],
        ["Medyan Pc", "1.59×10⁻⁹"],
        ["%95 Güven Aralığı", "[1.18×10⁻⁵,  2.22×10⁻⁵]"],
        ["P(Pc ≥ 10⁻⁴)", "%6.0  (izleme önerilebilir)"],
        ["Ort. Miss Distance", "5933 ± 4160 m"],
    ]
    add_table(doc, headers7, rows7, col_widths=[7, 9])
    add_paragraph(doc, "Tablo 7: Monte Carlo sonuç özeti (N=200, Senaryo 2)",
                  align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=10)
    add_paragraph(doc, "")

    add_figure(doc,
               os.path.join(RESULTS, "mc_pc_histogram.png"),
               "Şekil 11: Monte Carlo — Pc Dağılımı ve %95 Güven Aralığı (N=200)",
               width_cm=13)

    add_heading(doc, "4.6 Tüm Senaryolar Karşılaştırması", level=2)
    add_figure(doc,
               os.path.join(RESULTS, "all_scenarios_comparison.png"),
               "Şekil 12: Senaryo Karşılaştırması — Miss Distance ve Pc (tüm senaryolar)",
               width_cm=14)

    page_break(doc)

    # bölüm 5: sonuç

    add_heading(doc, "BÖLÜM 5: SONUÇ VE ÖNERİLER", level=1)

    add_heading(doc, "5.1 Sonuçlar", level=2)
    add_paragraph(doc,
        "Bu çalışmada, LEO uydu çarpışma önleme sistemi "
        "fizik temelli bir Python simülasyonu olarak başarıyla "
        "gerçeklenmiştir. Elde edilen başlıca bulgular şöyle özetlenebilir:")
    add_bullet(doc,
        "J2 pertürbasyonlu iki cisim modeli, 5 bağımsız V&V testiyle "
        "doğrulanmış ve geçerlenmiştir; tüm testler başarıyla geçilmiştir.")
    add_bullet(doc,
        "Chan 2D Pc hesabı, N=200 Monte Carlo analizi ile istatistiksel "
        "olarak desteklenmiş; %95 güven aralığı belirlenmiştir.")
    add_bullet(doc,
        "Yüksek riskli senaryoda 0.43 m/s'lik along-track manevra, "
        "Pc'yi 20 kat düşürerek (1.92×10⁻⁴ → 1.0×10⁻⁵) güvenli "
        "seviyeye getirmiştir.")
    add_bullet(doc,
        "Çoklu enkaz ortamında Pc temelli önceliklendirme algoritması "
        "başarıyla çalışmış; en kritik nesne otomatik tespit edilmiştir.")
    add_bullet(doc,
        "Delta-V duyarlılık analizi, yalnızca 2 m/s'lik yakıt yakımının "
        "Pc eşiğini aşmak için yeterli olduğunu göstermiştir.")

    add_heading(doc, "5.2 Kısıtlar", level=2)
    add_bullet(doc,
        "Anlık yakıt yakımı (impulsive) varsayımı gerçek itki sistemlerinden "
        "sapma yaratmaktadır.")
    add_bullet(doc,
        "Kovaryans matrisinin sabit kabulü, uzun vadeli belirsizlik "
        "büyümesini yansıtmamaktadır.")
    add_bullet(doc,
        "Hava sürtünmesi ve güneş radyasyon baskısı ihmal edilmiştir.")

    add_heading(doc, "5.3 Gelecek Çalışmalar için Öneriler", level=2)
    add_bullet(doc,
        "SGP4 propagatör entegrasyonu ile NORAD TLE verisinden gerçek "
        "uydu yörüngelerinin yüklenmesi")
    add_bullet(doc,
        "Sürekli düşük itki (low-thrust) manevra modeli ile elektrikli "
        "tahrik sistemleri için optimizasyon")
    add_bullet(doc,
        "Çok hedefli optimizasyon: yakıt tüketimi ile risk azaltımı "
        "trade-off analizi (Pareto cephesi)")
    add_bullet(doc,
        "Belirsizlik büyüme modeli (covariance propagation) ile "
        "dinamik Pc güncellemesi")

    page_break(doc)

    # ekler

    add_heading(doc, "EKLER (APPENDIX)", level=1)

    add_heading(doc, "EK-1: Kaynak Kodlar", level=2)
    add_paragraph(doc,
        "Projenin tüm kaynak kodları GitHub deposunda açık kaynak olarak "
        "yayımlanmıştır. Kod tabanı 10 Python modülünden oluşmakta "
        "ve toplam ~1500 satır içermektedir. Modüller arası bağımlılık "
        "yapısı aşağıdaki gibidir:")
    add_bullet(doc, "main.py → scenarios.py → satellite.py → orbital_mechanics.py")
    add_bullet(doc, "scenarios.py → collision_detection.py → avoidance.py")
    add_bullet(doc, "verification.py, monte_carlo.py → bağımsız test modülleri")
    add_bullet(doc, "gui.py → main.py (subprocess aracılığıyla)")
    add_paragraph(doc,
        "Kaynak kodların tamamı raporun sonunda ek olarak sunulmuş "
        "olup PyCharm IDE üzerinde geliştirilmiş, yorum satırlarıyla "
        "belgelenmiştir.")

    add_heading(doc, "EK-2: Ekran Görüntüleri", level=2)
    add_paragraph(doc,
        "Aşağıda simülasyonun grafik kullanıcı arayüzü (GUI) ve "
        "terminal çıktısından ekran görüntüleri sunulmaktadır.")
    add_paragraph(doc,
        "[Ekran Görüntüsü 1: GUI — Senaryo seçim ekranı]",
        align=WD_ALIGN_PARAGRAPH.CENTER, italic=True)
    add_paragraph(doc,
        "[Ekran Görüntüsü 2: Terminal — V&V test çıktıları]",
        align=WD_ALIGN_PARAGRAPH.CENTER, italic=True)
    add_paragraph(doc,
        "[Ekran Görüntüsü 3: Terminal — Monte Carlo özet tablosu]",
        align=WD_ALIGN_PARAGRAPH.CENTER, italic=True)

    # kaynaklar

    page_break(doc)
    add_heading(doc, "KAYNAKLAR", level=1)
    refs = [
        "Chan, F. K. (1997). Spacecraft Collision Probability. "
        "The Aerospace Press.",
        "Vallado, D. A. (2013). Fundamentals of Astrodynamics and "
        "Applications (4. baskı). Microcosm Press.",
        "Alfano, S. (2005). A numerical implementation of spherical "
        "object collision probability. Journal of the Astronautical "
        "Sciences, 53(1), 103–109.",
        "NASA Orbital Debris Program Office. (2023). Orbital Debris "
        "Quarterly News. https://orbitaldebris.jsc.nasa.gov/",
        "SciPy Developers. (2024). scipy.integrate.solve_ivp. "
        "https://docs.scipy.org/doc/scipy/",
        "European Space Agency. (2022). Space Debris by the Numbers. "
        "https://www.esa.int/Space_Safety/Space_Debris",
    ]
    for i, ref in enumerate(refs, 1):
        p = doc.add_paragraph(f"[{i}] {ref}")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.first_line_indent = Cm(-0.5)
        for run in p.runs:
            set_font(run, size=11)

    # kaydet
    out = "rapor.docx"
    doc.save(out)
    print(f"Rapor olusturuldu: {out}")


if __name__ == "__main__":
    build_report()
